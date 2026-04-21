import io
import re
from typing import Tuple
from fastapi import UploadFile, HTTPException
import pdfplumber
from docx import Document
from models.schemas import ParsedResume


SKILL_KEYWORDS = {
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
    "kotlin", "swift", "r", "scala", "ruby", "php", "dart",
    # Web
    "react", "angular", "vue", "nextjs", "nodejs", "express", "fastapi",
    "django", "flask", "spring", "graphql", "rest", "html", "css", "tailwind",
    # Data / ML
    "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "matplotlib",
    "spark", "hadoop", "sql", "postgresql", "mysql", "mongodb", "redis",
    "machine learning", "deep learning", "nlp", "computer vision", "llm",
    # Cloud / DevOps
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ci/cd",
    "github actions", "jenkins", "linux", "bash",
    # Tools
    "git", "jira", "figma", "postman", "elasticsearch",
}


async def extract_text_from_file(file: UploadFile) -> str:
    """Extract raw text from uploaded PDF or DOCX."""
    content = await file.read()
    filename = file.filename or ""

    if filename.lower().endswith(".pdf"):
        return _extract_pdf(content)
    elif filename.lower().endswith(".docx"):
        return _extract_docx(content)
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload PDF or DOCX."
        )


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                # Use layout-aware extraction to preserve spacing
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text:
                    text_parts.append(page_text)
        raw = "\n".join(text_parts).strip()

        # Fix common PDF spacing issues: add space before capital letters
        # that got merged (e.g. "DesignedandIntegrated" -> "Designed and Integrated")
        import re as _re
        # Fix missing spaces after punctuation
        raw = _re.sub(r'([a-z])([A-Z])', r'\1 \2', raw)
        # Fix merged words after commas with no space
        raw = _re.sub(r',([^\s])', r', \1', raw)
        # Fix multiple spaces
        raw = _re.sub(r' {2,}', ' ', raw)

        if not raw:
            raise HTTPException(status_code=422, detail="PDF appears to be empty or image-based (no extractable text).")
        return raw
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to parse PDF: {str(e)}")


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        raw = "\n".join(paragraphs).strip()
        if not raw:
            raise HTTPException(status_code=422, detail="DOCX appears to be empty.")
        return raw
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to parse DOCX: {str(e)}")


def parse_resume_sections(raw_text: str) -> ParsedResume:
    """Parse raw resume text into structured sections."""
    text_lower = raw_text.lower()

    # Extract skills via keyword matching
    found_skills = sorted([
        skill for skill in SKILL_KEYWORDS
        if skill in text_lower
    ])

    # Extract bullet points (lines starting with •, -, *, or numbered)
    bullet_pattern = re.compile(
        r"^[\s]*(?:[•\-\*\u2022\u2023\u25E6]|\d+[.)]) (.+)$",
        re.MULTILINE
    )
    bullets = [m.group(1).strip() for m in bullet_pattern.finditer(raw_text)]
    bullets = [b for b in bullets if len(b) > 20]  # filter trivial lines

    # Experience snippets (lines containing years or job titles)
    exp_pattern = re.compile(
        r"^.{10,}.{1}(20\d{2}|19\d{2}|present|current).{0,}$",
        re.MULTILINE | re.IGNORECASE
    )
    experience = list(set(
        m.group(0).strip() for m in exp_pattern.finditer(raw_text)
        if len(m.group(0).strip()) > 15
    ))[:10]

    # Project sections (look for "project" heading area)
    project_pattern = re.compile(
        r"(?:project[s]?|built|developed|created)[^\n]{5,}",
        re.IGNORECASE
    )
    projects = list(set(
        m.group(0).strip() for m in project_pattern.finditer(raw_text)
    ))[:8]

    return ParsedResume(
        raw_text=raw_text,
        skills=found_skills,
        experience_snippets=experience,
        projects=projects,
        bullet_points=bullets[:20],  # top 20 bullets for rewriting
    )