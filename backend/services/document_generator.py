"""
Document generation service for creating editable resume files.
Since preserving original PDF/DOCX formatting while applying edits is complex,
we provide the edited text in a clean, simple format that users can copy-paste
into their original resume.
"""
import io
from typing import List, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import HexColor
import re


def _parse_resume_structure(text: str) -> dict:
    """Parse resume text into structured sections."""
    sections = {
        "header": "",
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "other": []
    }
    
    lines = text.strip().split("\n")
    current_section = "header"
    current_content = []
    
    section_keywords = {
        "summary": ["summary", "objective", "profile", "about"],
        "experience": ["experience", "work history", "employment", "professional experience"],
        "education": ["education", "academic", "degree"],
        "skills": ["skills", "technologies", "technical skills", "competencies"],
        "projects": ["projects", "portfolio", "work samples"],
    }
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if this is a section header
        is_section_header = False
        for section_name, keywords in section_keywords.items():
            if any(keyword in line_lower for keyword in keywords) and len(line.strip()) < 50:
                if current_content:
                    if current_section == "header" and not sections["header"]:
                        sections["header"] = "\n".join(current_content)
                    elif current_section in ["experience", "education", "skills", "projects", "other"]:
                        sections[current_section].append("\n".join(current_content))
                    current_content = []
                current_section = section_name
                is_section_header = True
                break
        
        if not is_section_header:
            current_content.append(line)
    
    # Add remaining content
    if current_content:
        if current_section == "header" and not sections["header"]:
            sections["header"] = "\n".join(current_content)
        elif current_section in ["experience", "education", "skills", "projects", "other"]:
            sections[current_section].append("\n".join(current_content))
    
    return sections


def generate_docx(resume_text: str, filename: str = "resume_optimized.docx") -> Tuple[bytes, str]:
    """
    Generate a DOCX file from resume text.
    Uses simple formatting with clear structure for easy editing.
    """
    doc = Document()
    
    # Set document margins (standard resume margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process text line by line with intelligent formatting
    lines = resume_text.split("\n")
    
    # Track if we're in a bulleted section
    in_bullet_section = False
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            # Empty line - add spacing
            doc.add_paragraph()
            in_bullet_section = False
            continue
        
        # Check if line is a bullet point (starts with - • * or similar)
        is_bullet = any(line_stripped.startswith(char) for char in ['-', '•', '*', '◦', '○'])
        
        # Check if line looks like a section header
        # Headers are typically: ALL CAPS, or Title Case followed by colon, or very short
        is_header = (
            (len(line_stripped) < 40 and line_stripped.isupper()) or
            (line_stripped.endswith(':') and len(line_stripped) < 40) or
            (len(line_stripped) < 30 and line_stripped[0].isupper() and not any(c in line_stripped.lower() for c in ['the', 'and', 'with']))
        )
        
        if is_bullet:
            # Bullet point - add with bullet style
            # Remove the bullet character if present
            text = line_stripped.lstrip('-•*◦○ ')
            para = doc.add_paragraph(text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            in_bullet_section = True
            
        elif is_header:
            # Section header - make bold and larger
            if in_bullet_section:
                doc.add_paragraph()  # Add space before new section
            para = doc.add_paragraph(line_stripped)
            run = para.runs[0] if para.runs else para.add_run(line_stripped)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            in_bullet_section = False
            
        else:
            # Regular text - normal paragraph
            para = doc.add_paragraph(line_stripped)
            para.style = 'Normal'
            run = para.runs[0] if para.runs else None
            if run:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add small footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f"Optimized with ResumeIQ • {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue(), filename


def generate_pdf(resume_text: str, filename: str = "resume_optimized.pdf") -> Tuple[bytes, str]:
    """
    Generate a PDF file from resume text with clean, professional formatting.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=1*inch,
        bottomMargin=1*inch,
        leftMargin=1*inch,
        rightMargin=1*inch,
    )
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Header style (for section titles)
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=HexColor('#000000'),
        spaceAfter=6,
        spaceBefore=12,
        fontName='Helvetica-Bold',
        leading=14,
    )
    
    # Body style (for regular text)
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        textColor=HexColor('#000000'),
        spaceAfter=4,
        leading=13,
        fontName='Helvetica',
        alignment=TA_LEFT,
    )
    
    # Bullet style
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=20,
        bulletIndent=10,
        spaceAfter=4,
    )
    
    # Build document with intelligent formatting
    story = []
    lines = resume_text.split("\n")
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            # Empty line - add small spacing
            story.append(Spacer(1, 0.1*inch))
            continue
        
        # Escape XML special characters
        safe_line = (line_stripped
                     .replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;"))
        
        # Check if bullet point
        is_bullet = any(safe_line.startswith(char) for char in ['-', '•', '*', '◦', '○'])
        
        # Check if section header
        is_header = (
            (len(safe_line) < 40 and safe_line.isupper()) or
            (safe_line.endswith(':') and len(safe_line) < 40)
        )
        
        if is_bullet:
            # Bullet point - add with bullet
            text = safe_line.lstrip('-•*◦○ ')
            story.append(Paragraph(f"• {text}", bullet_style))
            
        elif is_header:
            # Section header - bold and larger
            story.append(Spacer(1, 0.05*inch))
            story.append(Paragraph(safe_line, header_style))
            
        else:
            # Regular text
            story.append(Paragraph(safe_line, body_style))
    
    # Add footer
    story.append(Spacer(1, 0.3*inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=HexColor('#888888'),
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique',
    )
    footer_text = f"<i>Optimized with ResumeIQ • {datetime.now().strftime('%B %d, %Y')}</i>"
    story.append(Paragraph(footer_text, footer_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer.getvalue(), filename


def apply_edits_to_text(original_text: str, suggestions: List[dict]) -> Tuple[str, str]:
    """
    Apply edit suggestions to the original resume text.
    Preserves original formatting and structure - only changes the specified text.
    Returns (edited_text, changes_summary).
    """
    edited_text = original_text
    changes = []
    
    for suggestion in suggestions:
        edit_type = suggestion.get("type", "replace")
        original = suggestion.get("original_text", "").strip()
        suggested = suggestion.get("suggested_text", "").strip()
        section = suggestion.get("section", "unknown")
        
        if not suggested:
            continue
        
        if edit_type == "replace" and original:
            # Replace exact match (case-sensitive first, then case-insensitive)
            if original in edited_text:
                edited_text = edited_text.replace(original, suggested, 1)
                changes.append(f"Replaced text in {section}")
            elif original.lower() in edited_text.lower():
                # Case-insensitive match
                import re
                pattern = re.compile(re.escape(original), re.IGNORECASE)
                edited_text = pattern.sub(suggested, edited_text, count=1)
                changes.append(f"Replaced text in {section}")
            else:
                # Fuzzy match: try matching key words
                original_words = original.split()
                if len(original_words) >= 3:
                    # Try to find a line containing most of the words
                    lines = edited_text.split("\n")
                    for i, line in enumerate(lines):
                        # Check if line contains at least 50% of the words
                        matches = sum(1 for word in original_words if word.lower() in line.lower())
                        if matches >= len(original_words) * 0.5:
                            lines[i] = suggested
                            edited_text = "\n".join(lines)
                            changes.append(f"Updated text in {section}")
                            break
        
        elif edit_type == "add":
            # Add text at the end of the document or after a section header
            section_headers = {
                "experience": ["experience", "work history", "employment"],
                "skills": ["skills", "technical skills", "competencies"],
                "education": ["education", "academic"],
                "projects": ["projects", "portfolio"],
                "summary": ["summary", "objective", "profile"],
            }
            
            added = False
            if section in section_headers:
                lines = edited_text.split("\n")
                for i, line in enumerate(lines):
                    line_lower = line.lower().strip()
                    if any(header in line_lower for header in section_headers[section]):
                        # Insert after this header line
                        # Find next empty line or next section
                        insert_pos = i + 1
                        while insert_pos < len(lines) and lines[insert_pos].strip():
                            insert_pos += 1
                        lines.insert(insert_pos, suggested)
                        edited_text = "\n".join(lines)
                        changes.append(f"Added content to {section}")
                        added = True
                        break
            
            if not added:
                # Append at the end
                edited_text += f"\n\n{suggested}"
                changes.append(f"Added content to {section}")
        
        elif edit_type == "remove" and original:
            # Remove exact match or line containing the text
            if original in edited_text:
                edited_text = edited_text.replace(original, "", 1)
                changes.append(f"Removed content from {section}")
            else:
                # Try to find and remove the line
                lines = edited_text.split("\n")
                for i, line in enumerate(lines):
                    if original.lower() in line.lower():
                        lines.pop(i)
                        edited_text = "\n".join(lines)
                        changes.append(f"Removed content from {section}")
                        break
        
        elif edit_type == "reword" and original:
            # Similar to replace
            if original in edited_text:
                edited_text = edited_text.replace(original, suggested, 1)
                changes.append(f"Reworded content in {section}")
    
    # Clean up multiple empty lines
    import re
    edited_text = re.sub(r'\n{3,}', '\n\n', edited_text)
    
    changes_summary = f"Applied {len(changes)} edits"
    if changes:
        changes_summary += ": " + "; ".join(changes[:3])
        if len(changes) > 3:
            changes_summary += f" and {len(changes) - 3} more"
    
    return edited_text, changes_summary
