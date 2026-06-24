"""
LaTeX service: Convert resumes to/from LaTeX format
This preserves ALL formatting and structure perfectly.
"""
import re
import subprocess
import tempfile
import os
import difflib
from pathlib import Path
from typing import Tuple
from groq import Groq
from core.config import settings


def get_client() -> Groq:
    return Groq(api_key=settings.GROQ_API_KEY)


async def resume_to_latex(resume_text: str, filename: str = "") -> str:
    """
    Convert plain resume text to LaTeX using DETERMINISTIC rule-based approach.
    This ensures consistent, compact formatting every time.
    """
    print("🔄 Using rule-based LaTeX generation for consistent formatting...")
    
    # Use deterministic approach instead of LLM
    return create_structured_latex(resume_text)


def create_structured_latex(resume_text: str) -> str:
    """
    Create structured LaTeX from resume text using deterministic rules.
    Preserves EXACT content and creates compact, 1-page formatting.
    """
    lines = resume_text.strip().split('\n')
    
    # Start with compact template
    latex = r"""\documentclass[11pt,a4paper]{article}
\usepackage[margin=0.5in,top=0.4in,bottom=0.4in]{geometry}
\usepackage{enumitem}
\usepackage{titlesec}
\usepackage{hyperref}
\usepackage[utf8]{inputenc}

\pagestyle{empty}
\setlength{\parindent}{0pt}
\setlist[itemize]{leftmargin=*,noitemsep,topsep=0pt,partopsep=0pt,parsep=0pt}
\setlength{\parskip}{0pt}

\titleformat{\section}{\large\bfseries}{}{0em}{}[\vspace{-8pt}\titlerule\vspace{2pt}]
\titlespacing{\section}{0pt}{6pt}{2pt}

\begin{document}

"""
    
    # Parse resume structure
    in_list = False
    first_line = True
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            if in_list:
                latex += "\\end{itemize}\n"
                in_list = False
            latex += "\n"
            continue
        
        # Detect name (usually first non-empty line)
        if first_line:
            latex += "\\begin{center}\n"
            latex += f"\\textbf{{\\Large {escape_latex(line_stripped)}}}\n"
            latex += "\\end{center}\n\n"
            first_line = False
            continue
        
        # Detect contact info (contains @ or phone numbers)
        if '@' in line_stripped or re.search(r'\+?\d{2,3}[-\s]?\d{3,4}[-\s]?\d{4}', line_stripped):
            latex += "\\begin{center}\n"
            latex += f"\\small {escape_latex(line_stripped)}\n"
            latex += "\\end{center}\n\n"
            continue
        
        # Detect section headers (short lines, typically all caps or title case, no bullet)
        is_section = (
            len(line_stripped) < 40 and
            not line_stripped.startswith(('-', '•', '*')) and
            (line_stripped.isupper() or (line_stripped[0].isupper() and line_stripped.count(' ') <= 3))
        )
        
        if is_section:
            if in_list:
                latex += "\\end{itemize}\n"
                in_list = False
            latex += f"\\section*{{{escape_latex(line_stripped)}}}\n"
            continue
        
        # Detect bullet points
        if line_stripped.startswith(('•', '-', '*', '◦', '○')):
            if not in_list:
                latex += "\\begin{itemize}\n"
                in_list = True
            # Remove bullet and clean
            text = line_stripped.lstrip('•-*◦○ ').strip()
            latex += f"\\item {escape_latex(text)}\n"
            continue
        
        # Regular text
        if in_list:
            latex += "\\end{itemize}\n"
            in_list = False
        
        # Check if it looks like a job title / company line (contains dates)
        if re.search(r'\d{4}', line_stripped):
            # Might contain dates - make it bold
            latex += f"\\textbf{{{escape_latex(line_stripped)}}}\n\n"
        else:
            latex += f"{escape_latex(line_stripped)}\n\n"
    
    if in_list:
        latex += "\\end{itemize}\n"
    
    latex += "\\end{document}"
    
    print(f"✅ Rule-based LaTeX generated ({len(latex)} chars) - deterministic formatting")
    return latex


def create_compact_latex_template(resume_text: str) -> str:
    """
    Fallback: Create a COMPACT LaTeX template for 1-page resumes.
    Simple wrapper that preserves exact text.
    """
    # Just use the structured approach
    return create_structured_latex(resume_text)


def escape_latex(text: str) -> str:
    """Escape special LaTeX characters."""
    replacements = {
        '\\': r'\textbackslash{}',
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


async def apply_edits_to_latex(latex_code: str, edits: list) -> Tuple[str, str]:
    """
    Apply edit suggestions to LaTeX code using MANUAL text replacement only.
    This preserves EXACT structure and formatting - only changes specific text content.
    """
    if not edits:
        return latex_code, "No edits to apply"
    
    print(f"📝 Applying {len(edits)} edits manually to preserve EXACT structure...")
    
    # Use ONLY manual approach - NO LLM to avoid structure changes
    return await _apply_edits_manually(latex_code, edits)


async def _apply_edits_manually(latex_code: str, edits: list) -> Tuple[str, str]:
    """
    Apply edits manually using smart text replacement.
    Preserves EXACT LaTeX structure - only changes the text content.
    """
    edited_latex = latex_code
    changes = []
    failed_edits = []
    
    for idx, edit in enumerate(edits, 1):
        edit_type = edit.get("type", "replace")
        original = edit.get("original_text", "").strip()
        suggested = edit.get("suggested_text", "").strip()
        section = edit.get("section", "unknown")
        
        if not suggested:
            continue
        
        print(f"  Edit {idx}: {edit_type} in {section}")
        print(f"    Original: {original[:80]}...")
        print(f"    Suggested: {suggested[:80]}...")
        
        # For replace/reword, we need to find and replace text
        if edit_type in ["replace", "reword"] and original:
            # Strategy 1: Find original text wrapped in LaTeX commands
            # Look for patterns like \textbf{original}, \textit{original}, or just original
            
            # Try multiple patterns to find the text
            patterns_to_try = [
                (original, suggested),  # Exact match
                (f"\\textbf{{{original}}}", f"\\textbf{{{suggested}}}"),
                (f"\\textit{{{original}}}", f"\\textit{{{suggested}}}"),
                (escape_latex(original), escape_latex(suggested)),
            ]
            
            replaced = False
            for orig_pattern, sugg_pattern in patterns_to_try:
                if orig_pattern in edited_latex:
                    edited_latex = edited_latex.replace(orig_pattern, sugg_pattern, 1)
                    changes.append(f"{edit_type.title()} in {section}")
                    print(f"    ✅ Applied using pattern: {orig_pattern[:50]}")
                    replaced = True
                    break
            
            if not replaced:
                # Strategy 2: Find the line containing the text and replace within that line
                lines = edited_latex.split('\n')
                for line_idx, line in enumerate(lines):
                    # Remove LaTeX commands to check if the text is present
                    clean_line = re.sub(r'\\[a-zA-Z]+\*?\{(.*?)\}', r'\1', line)
                    
                    if original.lower() in clean_line.lower():
                        # Found the line, now replace the text intelligently
                        # Replace within LaTeX commands if present
                        new_line = line
                        
                        # Try to replace within \textbf, \textit, etc.
                        if f"{{{original}}}" in line:
                            new_line = line.replace(f"{{{original}}}", f"{{{suggested}}}", 1)
                        elif original in line:
                            new_line = line.replace(original, suggested, 1)
                        
                        if new_line != line:
                            lines[line_idx] = new_line
                            edited_latex = '\n'.join(lines)
                            changes.append(f"{edit_type.title()} in {section}")
                            print(f"    ✅ Applied by line replacement")
                            replaced = True
                            break
            
            if not replaced:
                failed_edits.append(f"{edit_type} in {section}: '{original[:50]}...'")
                print(f"    ⚠️ Could not find text to replace")
        
        elif edit_type == "add":
            # Add new content after section header
            section_patterns = {
                "experience": r'\\section\*?\{.*?[Ee]xperience.*?\}',
                "skills": r'\\section\*?\{.*?[Ss]kills.*?\}',
                "education": r'\\section\*?\{.*?[Ee]ducation.*?\}',
                "projects": r'\\section\*?\{.*?[Pp]rojects.*?\}',
                "summary": r'\\section\*?\{.*?[Ss]ummary.*?\}',
            }
            
            pattern = section_patterns.get(section.lower())
            if pattern:
                match = re.search(pattern, edited_latex, re.IGNORECASE)
                if match:
                    insert_pos = match.end()
                    # Add with proper LaTeX formatting
                    edited_latex = (edited_latex[:insert_pos] + 
                                    f"\n{suggested}\n" + 
                                    edited_latex[insert_pos:])
                    changes.append(f"Added to {section}")
                    print(f"    ✅ Added content")
                else:
                    failed_edits.append(f"add in {section}: section not found")
                    print(f"    ⚠️ Section not found")
        
        elif edit_type == "remove" and original:
            # Remove text (keep structure)
            if original in edited_latex:
                edited_latex = edited_latex.replace(original, "", 1)
                changes.append(f"Removed from {section}")
                print(f"    ✅ Removed")
            else:
                failed_edits.append(f"remove in {section}")
                print(f"    ⚠️ Could not find text to remove")
    
    summary = f"Applied {len(changes)}/{len(edits)} edits - structure preserved"
    if failed_edits:
        print(f"  ⚠️ {len(failed_edits)} edits could not be applied:")
        for failed in failed_edits[:3]:
            print(f"    - {failed}")
    
    return edited_latex, summary


async def latex_to_pdf(latex_code: str) -> bytes:
    """
    Compile LaTeX code to PDF.
    Requires pdflatex to be installed on the system.
    """
    try:
        # Create temporary directory
        with tempfile.TemporaryDirectory() as tmpdir:
            tex_file = Path(tmpdir) / "resume.tex"
            pdf_file = Path(tmpdir) / "resume.pdf"
            
            # Write LaTeX code to file
            tex_file.write_text(latex_code, encoding='utf-8')
            
            # Compile with pdflatex
            result = subprocess.run(
                ['pdflatex', '-interaction=nonstopmode', '-output-directory', tmpdir, str(tex_file)],
                capture_output=True,
                timeout=30
            )
            
            if pdf_file.exists():
                pdf_bytes = pdf_file.read_bytes()
                print(f"✅ LaTeX compiled to PDF ({len(pdf_bytes)} bytes)")
                return pdf_bytes
            else:
                raise Exception("PDF not generated - pdflatex failed")
    
    except FileNotFoundError:
        raise Exception("pdflatex not installed. Please install TeX Live or MiKTeX.")
    except subprocess.TimeoutExpired:
        raise Exception("LaTeX compilation timeout")
    except Exception as e:
        raise Exception(f"LaTeX compilation failed: {str(e)}")


async def latex_to_docx_simple(latex_code: str) -> bytes:
    """
    Convert LaTeX to DOCX (simplified approach without pandoc).
    Extracts ALL text and creates a formatted DOCX with proper structure.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Extract text from LaTeX (remove commands but keep ALL content)
    text = latex_code
    
    # Remove preamble but keep everything after \begin{document}
    match = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', text, re.DOTALL)
    if match:
        text = match.group(1)
    
    # Process LaTeX commands while preserving content
    # Handle sections (make them headers)
    text = re.sub(r'\\section\*?\{(.*?)\}', r'\n\n###SECTION###\1###ENDSECTION###\n', text)
    text = re.sub(r'\\subsection\*?\{(.*?)\}', r'\n\n###SUBSECTION###\1###ENDSUBSECTION###\n', text)
    
    # Handle text formatting (keep the text, remove commands)
    text = re.sub(r'\\textbf\{(.*?)\}', r'###BOLD###\1###ENDBOLD###', text)
    text = re.sub(r'\\textit\{(.*?)\}', r'###ITALIC###\1###ENDITALIC###', text)
    text = re.sub(r'\\emph\{(.*?)\}', r'###ITALIC###\1###ENDITALIC###', text)
    
    # Handle lists
    text = re.sub(r'\\begin\{itemize\}', '###BEGINLIST###', text)
    text = re.sub(r'\\end\{itemize\}', '###ENDLIST###', text)
    text = re.sub(r'\\begin\{enumerate\}', '###BEGINLIST###', text)
    text = re.sub(r'\\end\{enumerate\}', '###ENDLIST###', text)
    text = re.sub(r'\\item\s+', '###BULLET### ', text)
    
    # Handle line breaks
    text = re.sub(r'\\\\', '\n', text)
    text = re.sub(r'\\newline', '\n', text)
    
    # Remove other LaTeX commands while preserving their content
    text = re.sub(r'\\[a-zA-Z]+\*?\[.*?\]\{(.*?)\}', r'\1', text)  # Commands with optional args
    text = re.sub(r'\\[a-zA-Z]+\*?\{(.*?)\}', r'\1', text)  # Commands with content
    text = re.sub(r'\\[a-zA-Z]+\*?', '', text)  # Commands without content
    
    # Clean up special characters
    text = text.replace('\\&', '&')
    text = text.replace('\\%', '%')
    text = text.replace('\\$', '$')
    text = text.replace('\\#', '#')
    text = text.replace('\\_', '_')
    text = text.replace('\\{', '{')
    text = text.replace('\\}', '}')
    text = text.replace('~', ' ')
    
    # Remove multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Create DOCX with proper formatting
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Process the text line by line with markers
    lines = text.split('\n')
    in_list = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            if not in_list:
                doc.add_paragraph()
            continue
        
        # Handle list markers
        if '###BEGINLIST###' in line:
            in_list = True
            continue
        elif '###ENDLIST###' in line:
            in_list = False
            continue
        
        # Handle bullet points
        if '###BULLET###' in line:
            line = line.replace('###BULLET###', '')
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            _add_formatted_text(para, line)
            continue
        
        # Handle section headers
        if '###SECTION###' in line:
            line = line.replace('###SECTION###', '').replace('###ENDSECTION###', '')
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(line)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            continue
        
        # Handle subsection headers
        if '###SUBSECTION###' in line:
            line = line.replace('###SUBSECTION###', '').replace('###ENDSUBSECTION###', '')
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(line)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            continue
        
        # Regular paragraph
        para = doc.add_paragraph()
        _add_formatted_text(para, line)
    
    # Save to bytes
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_formatted_text(paragraph, text: str):
    """Add text to paragraph with inline formatting markers."""
    from docx.shared import Pt, RGBColor
    
    # Split by formatting markers
    parts = re.split(r'(###BOLD###|###ENDBOLD###|###ITALIC###|###ENDITALIC###)', text)
    
    bold = False
    italic = False
    
    for part in parts:
        if part == '###BOLD###':
            bold = True
        elif part == '###ENDBOLD###':
            bold = False
        elif part == '###ITALIC###':
            italic = True
        elif part == '###ENDITALIC###':
            italic = False
        elif part:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
